import re
from typing import List, Dict, Tuple, Optional, Callable
from dataclasses import dataclass
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_LINE_SPACING

from .config import config
from .fonts import get_preferred_font, AVAILABLE_FONTS

@dataclass
class CheckResult:
    level: str
    category: str
    message: str
    location: Optional[str] = None
    suggestion: Optional[str] = None

class FormatChecker:
    LEVEL_ERROR = "error"
    LEVEL_WARNING = "warning"
    LEVEL_INFO = "info"

    CATEGORY_PAGE = "页面设置"
    CATEGORY_FONT = "字体规范"
    CATEGORY_HEADING = "标题层级"
    CATEGORY_BODY = "正文格式"
    CATEGORY_LINE_SPACING = "行距"
    CATEGORY_SIGNATURE = "落款"
    CATEGORY_TABLE = "表格"
    CATEGORY_PAGINATION = "页码"

    def __init__(self, progress_callback: Optional[Callable[[str, int], None]] = None):
        self.progress_callback = progress_callback
        self.results: List[CheckResult] = []
        self._fonts = None

    def _get_fonts(self):
        if self._fonts is None:
            cfg_fonts = config.get_fonts()
            self._fonts = {}
            for key, font_name in cfg_fonts.items():
                if font_name in AVAILABLE_FONTS.get(key, []):
                    self._fonts[key] = font_name
                else:
                    self._fonts[key] = get_preferred_font(key)
        return self._fonts

    def _report(self, level: str, category: str, message: str,
                location: Optional[str] = None, suggestion: Optional[str] = None):
        self.results.append(CheckResult(level, category, message, location, suggestion))

    def check_document(self, file_path: str) -> List[CheckResult]:
        self.results = []

        try:
            if self.progress_callback:
                self.progress_callback("正在加载文档...", 5)
            
            from utils import logger
            logger.debug(f"Checker: 开始加载文档 {file_path}")
            doc = Document(file_path)
            logger.debug("Checker: 文档加载完成")

            if self.progress_callback:
                self.progress_callback("正在检查页面设置...", 10)
            
            logger.debug("Checker: 开始检查页面设置")
            self._check_page_setup(doc)
            logger.debug("Checker: 页面设置检查完成")
            
            self._check_paragraphs(doc)
            logger.debug("Checker: 段落检查完成")

            if self.progress_callback:
                self.progress_callback("正在检查字体规范...", 40)

            self._check_fonts(doc)
            logger.debug("Checker: 字体检查完成")

            if self.progress_callback:
                self.progress_callback("正在检查标题层级...", 60)

            self._check_headings(doc)
            logger.debug("Checker: 标题层级检查完成")

            if self.progress_callback:
                self.progress_callback("正在检查表格...", 80)

            self._check_tables(doc)
            logger.debug("Checker: 表格检查完成")

            if self.progress_callback:
                self.progress_callback(f"检查完成: {len(self.results)} 个问题", 100)

            return self.results
        except Exception as e:
            from utils import logger
            logger.error(f"Checker.check_document 异常: {e}", exc_info=True)
            raise e

    def _check_page_setup(self, doc: Document):
        if not doc.sections:
            return
            
        section = doc.sections[0]
        expected_margins = config.get_margins()

        # python-docx properties can sometimes be None
        try:
            page_width_emu = section.page_width
            page_height_emu = section.page_height
            
            if page_width_emu is None or page_height_emu is None:
                page_width = 21.0
                page_height = 29.7
            else:
                page_width = page_width_emu / 914400
                page_height = page_height_emu / 914400
        except Exception:
            page_width = 21.0
            page_height = 29.7

        if abs(page_width - 21.0) > 0.1 or abs(page_height - 29.7) > 0.1:
            self._report(
                self.LEVEL_ERROR,
                self.CATEGORY_PAGE,
                f"纸张尺寸异常: {page_width:.1f}cm x {page_height:.1f}cm，应为 A4 (21cm x 29.7cm)",
                suggestion="使用标准 A4 纸张"
            )

        def get_margin_mm(emu_val):
            if emu_val is None:
                return 0.0
            return emu_val / 914400 * 10

        top_margin = get_margin_mm(section.top_margin)
        bottom_margin = get_margin_mm(section.bottom_margin)
        left_margin = get_margin_mm(section.left_margin)
        right_margin = get_margin_mm(section.right_margin)

        tolerance = 2
        if abs(top_margin - expected_margins["top"]) > tolerance:
            self._report(
                self.LEVEL_ERROR,
                self.CATEGORY_PAGE,
                f"上边距: {top_margin:.1f}mm，标准为 {expected_margins['top']}mm",
                suggestion=f"建议设置为 {expected_margins['top']}mm"
            )

        if abs(bottom_margin - expected_margins["bottom"]) > tolerance:
            self._report(
                self.LEVEL_ERROR,
                self.CATEGORY_PAGE,
                f"下边距: {bottom_margin:.1f}mm，标准为 {expected_margins['bottom']}mm",
                suggestion=f"建议设置为 {expected_margins['bottom']}mm"
            )

        if abs(left_margin - expected_margins["left"]) > tolerance:
            self._report(
                self.LEVEL_ERROR,
                self.CATEGORY_PAGE,
                f"左边距: {left_margin:.1f}mm，标准为 {expected_margins['left']}mm",
                suggestion=f"建议设置为 {expected_margins['left']}mm"
            )

        if abs(right_margin - expected_margins["right"]) > tolerance:
            self._report(
                self.LEVEL_ERROR,
                self.CATEGORY_PAGE,
                f"右边距: {right_margin:.1f}mm，标准为 {expected_margins['right']}mm",
                suggestion=f"建议设置为 {expected_margins['right']}mm"
            )

    def _get_line_spacing(self, para):
        pf = para.paragraph_format
        if pf.line_spacing is None:
            return 28.0 # Default fallback
        return pf.line_spacing / 20

    def _get_font_size(self, run):
        if run.font.size is None:
            return 12.0 # Default fallback
        return run.font.size.pt

    def _get_bold(self, run):
        if run.font.bold is None:
            return False # Default fallback
        return run.font.bold

    def _check_paragraphs(self, doc: Document):
        expected_line_spacing = config.get_line_spacing()
        fonts = self._get_fonts()

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            location = f"第 {idx + 1} 段"

            pf = para.paragraph_format
            if pf.line_spacing_rule == WD_LINE_SPACING.EXACTLY:
                actual_spacing = self._get_line_spacing(para)
                if abs(actual_spacing - expected_line_spacing) > 2:
                    self._report(
                        self.LEVEL_WARNING,
                        self.CATEGORY_LINE_SPACING,
                        f"{location}: 行距为 {actual_spacing:.0f}pt，标准为固定值 {expected_line_spacing}pt",
                        location,
                        f"设置行距为固定值 {expected_line_spacing}pt"
                    )

            for run in para.runs:
                if not run.text.strip():
                    continue

                is_bold = self._get_bold(run)
                text_content = run.text

                heading_patterns = [
                    (r'^[一二三四五六七八九十]+、', 1, '一级标题'),
                    (r'^（[一二三四五六七八九十]+）', 2, '二级标题'),
                    (r'^\d+\.', 3, '三级标题'),
                    (r'^（\d+）', 4, '四级标题'),
                ]

                is_heading = False
                for pattern, level, heading_name in heading_patterns:
                    if re.match(pattern, text_content):
                        is_heading = True
                        break

                if not is_heading and is_bold and text_content:
                    self._report(
                        self.LEVEL_ERROR,
                        self.CATEGORY_BODY,
                        f"{location}: 正文中不应使用加粗格式",
                        location,
                        "去除正文中的加粗格式"
                    )
                    break

    def _check_fonts(self, doc: Document):
        fonts = self._get_fonts()
        font_sizes = config.get_font_sizes()

        expected_fonts = {
            'title': fonts.get('title', fonts.get('body')),
            'heading1': fonts.get('heading1', fonts.get('body')),
            'heading2': fonts.get('heading2', fonts.get('body')),
            'heading3': fonts.get('heading3', fonts.get('body')),
            'heading4': fonts.get('heading4', fonts.get('body')),
            'body': fonts.get('body'),
        }

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            if not para.runs:
                continue

            first_run = para.runs[0]
            location = f"第 {idx + 1} 段"

            ptype = self._classify_paragraph(para)

            if ptype == 'title':
                actual_size = self._get_font_size(first_run)
                if abs(actual_size - font_sizes['title']) > 1:
                    self._report(
                        self.LEVEL_ERROR,
                        self.CATEGORY_FONT,
                        f"{location}: 主标题字号应为 {font_sizes['title']}pt",
                        location,
                        f"设置主标题为 {font_sizes['title']}pt"
                    )
                if para.alignment != 1:
                    self._report(
                        self.LEVEL_WARNING,
                        self.CATEGORY_FONT,
                        f"{location}: 主标题应居中对齐",
                        location,
                        "设置标题居中对齐"
                    )

            elif ptype in ('heading1', 'heading2', 'heading3', 'heading4'):
                actual_size = self._get_font_size(first_run)
                if abs(actual_size - font_sizes['heading']) > 1:
                    self._report(
                        self.LEVEL_ERROR,
                        self.CATEGORY_FONT,
                        f"{location}: 标题字号应为 {font_sizes['heading']}pt",
                        location,
                        f"设置标题为 {font_sizes['heading']}pt"
                    )

                expected_bold = ptype in ('heading1', 'heading2', 'heading3')
                actual_bold = self._get_bold(first_run)

                if actual_bold != expected_bold:
                    level_name = {'heading1': '一、', 'heading2': '（一）', 'heading3': '1.', 'heading4': '（1）'}
                    if expected_bold:
                        self._report(
                            self.LEVEL_ERROR,
                            self.CATEGORY_FONT,
                            f"{location}: {level_name.get(ptype, '')}应加粗",
                            location,
                            "添加加粗格式"
                        )
                    else:
                        self._report(
                            self.LEVEL_WARNING,
                            self.CATEGORY_FONT,
                            f"{location}: {level_name.get(ptype, '')}不应加粗",
                            location,
                            "去除加粗格式"
                        )

    def _classify_paragraph(self, para) -> str:
        text = para.text.strip()
        if not text:
            return 'body'

        if text.startswith('#'):
            return 'title'

        heading_patterns = [
            (r'^[一二三四五六七八九十]+、', 'heading1'),
            (r'^（[一二三四五六七八九十]+）', 'heading2'),
            (r'^\d+\.', 'heading3'),
            (r'^（\d+）', 'heading4'),
        ]

        for pattern, ptype in heading_patterns:
            if re.match(pattern, text):
                return ptype

        return 'body'

    def _check_headings(self, doc: Document):
        heading_order = []
        current_level = 0

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            ptype = self._classify_paragraph(para)
            if ptype.startswith('heading'):
                level_map = {'heading1': 1, 'heading2': 2, 'heading3': 3, 'heading4': 4}
                level = level_map.get(ptype, 0)

                heading_order.append((idx + 1, level, text[:20]))

                if level > current_level + 1 and current_level > 0:
                    self._report(
                        self.LEVEL_ERROR,
                        self.CATEGORY_HEADING,
                        f"第 {idx + 1} 段: 标题层级跳跃（从第 {current_level} 级跳到第 {level} 级）",
                        f"第 {idx + 1} 段",
                        "检查标题层级顺序"
                    )

                current_level = level

    def _check_tables(self, doc: Document):
        font_sizes = config.get_font_sizes()

        for idx, table in enumerate(doc.tables):
            location = f"表格 {idx + 1}"

            if table.rows:
                header_row = table.rows[0]
                for cell in header_row.cells:
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        first_run = cell.paragraphs[0].runs[0]
                        if not self._get_bold(first_run):
                            self._report(
                                self.LEVEL_WARNING,
                                self.CATEGORY_TABLE,
                                f"{location} 表头未加粗",
                                location,
                                "设置表头加粗"
                            )

                        actual_size = self._get_font_size(first_run)
                        if abs(actual_size - font_sizes['table']) > 1:
                            self._report(
                                self.LEVEL_WARNING,
                                self.CATEGORY_TABLE,
                                f"{location} 表头字号应为 {font_sizes['table']}pt",
                                location,
                                f"设置表头为 {font_sizes['table']}pt"
                            )

    def get_summary(self) -> Dict[str, int]:
        summary = {
            'error': 0,
            'warning': 0,
            'info': 0,
            'total': len(self.results)
        }

        for result in self.results:
            if result.level == self.LEVEL_ERROR:
                summary['error'] += 1
            elif result.level == self.LEVEL_WARNING:
                summary['warning'] += 1
            elif result.level == self.LEVEL_INFO:
                summary['info'] += 1

        return summary

    def generate_report(self) -> str:
        summary = self.get_summary()

        report_lines = [
            "=" * 50,
            "格式检查报告",
            "=" * 50,
            f"总计: {summary['total']} 个问题",
            f"  - 错误: {summary['error']}",
            f"  - 警告: {summary['warning']}",
            f"  - 提示: {summary['info']}",
            "=" * 50,
            ""
        ]

        if not self.results:
            report_lines.append("✅ 未发现问题，文档格式符合标准")
            return "\n".join(report_lines)

        current_category = None
        for result in self.results:
            if result.category != current_category:
                current_category = result.category
                report_lines.append(f"\n【{current_category}】")

            level_icon = {
                self.LEVEL_ERROR: "❌",
                self.LEVEL_WARNING: "⚠️",
                self.LEVEL_INFO: "ℹ️"
            }.get(result.level, "•")

            line = f"  {level_icon} {result.message}"
            if result.location:
                line += f" ({result.location})"
            report_lines.append(line)

            if result.suggestion:
                report_lines.append(f"     → 建议: {result.suggestion}")

        return "\n".join(report_lines)
