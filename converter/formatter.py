import re
from pathlib import Path
from typing import Optional, Callable
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .config import config
from .fonts import get_preferred_font, AVAILABLE_FONTS

class FormatFormatter:
    def __init__(self, progress_callback: Optional[Callable[[str, int], None]] = None):
        self.progress_callback = progress_callback
        self._fonts = None

    def _get_fonts(self):
        if self._fonts is None:
            cfg_fonts = config.get_fonts()
            self._fonts = {}
            for key, font_name in cfg_fonts.items():
                if font_name in AVAILABLE_FONTS.get(key, []):
                    self._fonts[key] = font_name
                else:
                    self._fonts[key] = get_preferred_font(key)
        return self._fonts

    def _set_font(self, run, font_name, size, bold=False):
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = False
        run.font.underline = False
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.strike = False
        run.font.subscript = False
        run.font.superscript = False

        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)

    def _apply_paragraph_format(self, para, font_name, font_size, bold=False,
                                align='left', indent=0, line_spacing=None):
        if line_spacing is None:
            line_spacing = config.get_line_spacing()

        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)

        pf = para.paragraph_format
        pf.left_indent = Pt(0)
        pf.right_indent = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_spacing)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        if indent > 0:
            pf.first_line_indent = Pt(indent)
        else:
            pf.first_line_indent = Pt(0)

        for run in para.runs:
            self._set_font(run, font_name, font_size, bold)

        return para

    def _set_page_setup(self, doc):
        margins = config.get_margins()
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(margins["top"] / 10)
        section.bottom_margin = Cm(margins["bottom"] / 10)
        section.left_margin = Cm(margins["left"] / 10)
        section.right_margin = Cm(margins["right"] / 10)

    def _classify_paragraph(self, para, is_first_non_empty=False) -> str:
        text = para.text.strip()
        if not text:
            return 'body'

        # Markdown style title
        if text.startswith('#'):
            return 'title'

        # If it's the first non-empty paragraph and not a heading, it's likely the title
        heading_patterns = [
            (r'^[一二三四五六七八九十]+、', 'heading1'),
            (r'^（[一二三四五六七八九十]+）', 'heading2'),
            (r'^\d+\.', 'heading3'),
            (r'^（\d+）', 'heading4'),
        ]

        is_heading = False
        for pattern, ptype in heading_patterns:
            if re.match(pattern, text):
                is_heading = True
                if not is_first_non_empty: # Don't return title if it's clearly a heading
                    return ptype
                break

        if is_first_non_empty and not is_heading:
            return 'title'
        
        if is_heading:
            # Re-run matching to return the correct ptype
            for pattern, ptype in heading_patterns:
                if re.match(pattern, text):
                    return ptype

        return 'body'

    def format_document(self, input_path: str, output_path: str) -> bool:
        try:
            if self.progress_callback:
                self.progress_callback("正在加载文档...", 10)

            doc = Document(input_path)

            if self.progress_callback:
                self.progress_callback("正在设置页面格式...", 20)

            self._set_page_setup(doc)

            if self.progress_callback:
                self.progress_callback("正在格式化段落...", 40)

            fonts = self._get_fonts()
            font_sizes = config.get_font_sizes()
            line_spacing = config.get_line_spacing()
            char_indent = config.get_char_indent()

            title_found = False
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Pass title_found to help classify the first non-empty paragraph as title
                ptype = self._classify_paragraph(para, is_first_non_empty=not title_found)
                
                if ptype == 'title':
                    title_found = True
                    # Clean markdown # if present
                    para.text = para.text.lstrip('#').strip()
                    self._apply_paragraph_format(
                        para, fonts['title'], font_sizes['title'],
                        bold=False, align='center',
                        indent=0, line_spacing=line_spacing
                    )
                elif ptype == 'heading1':
                    self._apply_paragraph_format(
                        para, fonts['heading1'], font_sizes['heading'],
                        bold=True, align='left',
                        indent=char_indent, line_spacing=line_spacing
                    )
                elif ptype == 'heading2':
                    self._apply_paragraph_format(
                        para, fonts['heading2'], font_sizes['heading'],
                        bold=True, align='left',
                        indent=char_indent, line_spacing=line_spacing
                    )
                elif ptype == 'heading3':
                    self._apply_paragraph_format(
                        para, fonts['heading3'], font_sizes['heading'],
                        bold=True, align='left',
                        indent=char_indent, line_spacing=line_spacing
                    )
                elif ptype == 'heading4':
                    self._apply_paragraph_format(
                        para, fonts['heading4'], font_sizes['heading'],
                        bold=False, align='left',
                        indent=char_indent, line_spacing=line_spacing
                    )
                else:
                    self._apply_paragraph_format(
                        para, fonts['body'], font_sizes['body'],
                        bold=False, align='justify',
                        indent=char_indent, line_spacing=line_spacing
                    )

                    for run in para.runs:
                        run.font.bold = False

            if self.progress_callback:
                self.progress_callback("正在格式化表格...", 70)

            self._format_tables(doc)

            if self.progress_callback:
                self.progress_callback("正在设置页码...", 90)

            self._set_page_numbers(doc)

            if self.progress_callback:
                self.progress_callback("正在保存文档...", 95)

            doc.save(output_path)

            if self.progress_callback:
                self.progress_callback(f"格式化完成: {Path(output_path).name}", 100)

            return True

        except Exception as e:
            if self.progress_callback:
                self.progress_callback(f"格式化失败: {str(e)}", 0)
            return False

    def _format_tables(self, doc: Document):
        font_sizes = config.get_font_sizes()
        fonts = self._get_fonts()

        for table in doc.tables:
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            tblBorders = OxmlElement('w:tblBorders')

            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tblBorders.append(border)

            existing_borders = tblPr.find(qn('w:tblBorders'))
            if existing_borders is not None:
                tblPr.remove(existing_borders)
            tblPr.append(tblBorders)

            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                        para.paragraph_format.line_spacing = Pt(16)
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)

                        for run in para.runs:
                            if row_idx == 0:
                                self._set_font(run, fonts['table'], font_sizes['table'], bold=True)
                            else:
                                self._set_font(run, fonts['table'], font_sizes['table'], bold=False)

    def _set_page_numbers(self, doc: Document):
        section = doc.sections[0]

        try:
            doc.settings.odd_and_even_pages_header_footer = True
        except Exception:
            settings_el = doc.settings._element
            if settings_el.find(qn('w:evenAndOddHeaders')) is None:
                settings_el.append(OxmlElement('w:evenAndOddHeaders'))

        section.odd_and_even_pages_header_footer = True
        section.footer_distance = Cm(0.7)

        odd_footer = section.footer
        even_footer = section.even_page_footer
        odd_footer.is_linked_to_previous = False
        even_footer.is_linked_to_previous = False

        for para in odd_footer.paragraphs:
            para.clear()
        for para in even_footer.paragraphs:
            para.clear()

        def build_footer_line(footer, align, pad_fullwidth):
            if footer.paragraphs:
                para = footer.paragraphs[0]
            else:
                para = footer.add_paragraph()

            para.alignment = align

            if pad_fullwidth:
                run0 = para.add_run("　")
                self._set_font(run0, 'Times New Roman', 14, bold=False)

            run1 = para.add_run("— ")
            self._set_font(run1, 'Times New Roman', 14, bold=False)

            run2 = para.add_run()
            self._set_font(run2, 'Times New Roman', 14, bold=False)
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run2._r.append(fldChar1)

            run3 = para.add_run()
            self._set_font(run3, 'Times New Roman', 14, bold=False)
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            run3._r.append(instrText)

            run4 = para.add_run()
            self._set_font(run4, 'Times New Roman', 14, bold=False)
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            run4._r.append(fldChar2)

            run5 = para.add_run()
            self._set_font(run5, 'Times New Roman', 14, bold=False)
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')
            run5._r.append(fldChar3)

            run6 = para.add_run(" —")
            self._set_font(run6, 'Times New Roman', 14, bold=False)

            if not pad_fullwidth:
                run7 = para.add_run("　")
                self._set_font(run7, 'Times New Roman', 14, bold=False)

        build_footer_line(odd_footer, WD_ALIGN_PARAGRAPH.RIGHT, pad_fullwidth=True)
        build_footer_line(even_footer, WD_ALIGN_PARAGRAPH.LEFT, pad_fullwidth=False)
