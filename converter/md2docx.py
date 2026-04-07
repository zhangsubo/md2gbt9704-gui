import re
from pathlib import Path
from typing import Optional, Callable
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .config import config
from .fonts import get_preferred_font, AVAILABLE_FONTS

class MarkdownConverter:
    def __init__(self, progress_callback: Optional[Callable[[str, int], None]] = None):
        self.progress_callback = progress_callback
        self._fonts = None

    def _get_fonts(self):
        if self._fonts is None:
            cfg_fonts = config.get_fonts()
            self._fonts = {}
            for key, font_name in cfg_fonts.items():
                if font_name in AVAILABLE_FONTS.get(key, []):
                    self._fonts[key] = font_name
                else:
                    self._fonts[key] = get_preferred_font(key)
        return self._fonts

    def _set_font(self, run, font_name, size, bold=False):
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = False
        run.font.underline = False
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.strike = False
        run.font.subscript = False
        run.font.superscript = False
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)

    def _set_page_setup(self, doc):
        margins = config.get_margins()
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(margins["top"] / 10)
        section.bottom_margin = Cm(margins["bottom"] / 10)
        section.left_margin = Cm(margins["left"] / 10)
        section.right_margin = Cm(margins["right"] / 10)

    def _apply_paragraph_style(self, para, font_name, font_size, bold=False,
                              align='left', indent=0, line_spacing=None):
        if line_spacing is None:
            line_spacing = config.get_line_spacing()
        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        pf = para.paragraph_format
        pf.left_indent = Pt(0)
        pf.right_indent = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_spacing)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        if indent > 0:
            pf.first_line_indent = Pt(indent)
        else:
            pf.first_line_indent = Pt(0)
        for run in para.runs:
            self._set_font(run, font_name, font_size, bold)
        return para

    def _add_empty_paragraph(self, doc, line_spacing=None):
        if line_spacing is None:
            line_spacing = config.get_line_spacing()
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        para.paragraph_format.line_spacing = Pt(line_spacing)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run('')
        fonts = self._get_fonts()
        self._set_font(run, fonts['body'], config.get_font_sizes()['body'], bold=False)
        return para

    def _clean_markdown_text(self, text):
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'^[\s]*[\*\+\-][\s]+', '', text)
        text = re.sub(r'^[\s]*\d+\.[\s]+', '', text)
        return text

    def _normalize_spacing(self, text):
        text = re.sub(r'([、，。！？；：])\s+', r'\1', text)
        text = re.sub(r'([\u4e00-\u9fff])\s+([\u4e00-\u9fff])', r'\1\2', text)
        text = re.sub(r'([\u4e00-\u9fff])([a-zA-Z0-9])', r'\1 \2', text)
        text = re.sub(r'([a-zA-Z0-9])([\u4e00-\u9fff])', r'\1 \2', text)
        return text

    def _is_signature_line(self, text):
        if re.match(r'^\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日$', text.strip()):
            return 'date'
        if re.match(r'^.{2,}(公司|局|委|处|部|办|中心|厅|院|所|站|组|会|署|集团|有限公司)$', text.strip()):
            return 'organization'
        return None

    def _parse_markdown(self, content):
        lines = content.split('\n')
        result = []
        i = 0

        while i < len(lines):
            stripped = lines[i].strip()
            if not stripped:
                i += 1
                continue

            if '|' in stripped and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if re.match(r'^\|[\s\-:]+\|', next_line):
                    table_rows = []
                    header_cells = [cell.strip() for cell in stripped.split('|')[1:-1]]
                    table_rows.append(header_cells)
                    i += 2
                    while i < len(lines):
                        row_line = lines[i].strip()
                        if '|' in row_line and row_line.startswith('|'):
                            row_cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                            table_rows.append(row_cells)
                            i += 1
                        else:
                            break
                    result.append((None, '', 'table', table_rows))
                    continue

            if stripped.startswith('# ') and len(stripped) > 2:
                result.append((0, stripped[2:].strip(), 'title', None))
                i += 1
                continue

            text = stripped
            for prefix in ['#### ', '### ', '## ', '# ']:
                if text.startswith(prefix):
                    text = text[len(prefix):].strip()
                    break

            clean_text = self._clean_markdown_text(text)
            sig_type = self._is_signature_line(clean_text)
            if sig_type:
                result.append((None, clean_text, 'signature', sig_type))
                i += 1
                continue

            if re.match(r'^[一二三四五六七八九十]+、', clean_text):
                result.append((1, clean_text, 'heading1', None))
            elif re.match(r'^（[一二三四五六七八九十]+）', clean_text) or re.match(r'^\([一二三四五六七八九十]+\)', clean_text):
                result.append((2, clean_text, 'heading2', None))
            elif re.match(r'^\d+\.', clean_text):
                result.append((3, clean_text, 'heading3', None))
            elif re.match(r'^（\d+）', clean_text) or re.match(r'^\(\d+\)', clean_text):
                result.append((4, clean_text, 'heading4', None))
            else:
                result.append((None, clean_text, 'body', None))
            i += 1
        return result

    def _create_table(self, doc, table_data):
        if not table_data:
            return
        fonts = self._get_fonts()
        font_size = config.get_font_sizes()['table']
        font_name = fonts['table']

        rows = len(table_data)
        cols = len(table_data[0]) if table_data else 0
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        tblBorders = OxmlElement('w:tblBorders')

        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)

        tblPr.append(tblBorders)

        for row_idx, row_data in enumerate(table_data):
            row = table.rows[row_idx]
            for col_idx, cell_text in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.text = ''
                clean_text = self._clean_markdown_text(cell_text.strip())
                clean_text = self._normalize_spacing(clean_text)
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(clean_text)
                if row_idx == 0:
                    self._set_font(run, font_name, font_size, bold=True)
                else:
                    self._set_font(run, font_name, font_size, bold=False)
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                para.paragraph_format.line_spacing = Pt(16)
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)

        self._add_empty_paragraph(doc)
        return table

    def _create_signature_paragraph(self, doc, text, sig_type):
        fonts = self._get_fonts()
        font_size = config.get_font_sizes()['body']
        para = doc.add_paragraph()
        run = para.add_run(text)
        self._set_font(run, fonts['body'], font_size, bold=False)
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf = para.paragraph_format
        pf.left_indent = Pt(0)
        pf.right_indent = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(config.get_line_spacing())
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        return para

    def _create_styled_paragraph(self, doc, text, ptype):
        fonts = self._get_fonts()
        font_sizes = config.get_font_sizes()
        line_spacing = config.get_line_spacing()
        char_indent = config.get_char_indent()

        clean_text = self._normalize_spacing(text)

        prefix_patterns = [
            r'^([一二三四五六七八九十]+、)(.*)',
            r'^(（[一二三四五六七八九十]+）)(.*)',
            r'^(\([一二三四五六七八九十]+\))(.*)',
            r'^(\d+\.)(.*)',
            r'^(（\d+）)(.*)',
            r'^(\(\d+\))(.*)',
        ]

        for pattern in prefix_patterns:
            match = re.match(pattern, clean_text)
            if match:
                prefix, rest = match.groups()
                rest = rest.lstrip()
                clean_text = prefix + rest
                break

        para = doc.add_paragraph()
        run = para.add_run(clean_text)

        if ptype == 'title':
            font_name = fonts['title']
            self._apply_paragraph_style(para, font_name, font_sizes["title"],
                                        bold=False, align='center',
                                        indent=0, line_spacing=line_spacing)
        elif ptype == 'heading1':
            font_name = fonts['heading1']
            self._apply_paragraph_style(para, font_name, font_sizes["heading"],
                                        bold=True, align='left',
                                        indent=char_indent, line_spacing=line_spacing)
        elif ptype == 'heading2':
            font_name = fonts['heading2']
            self._apply_paragraph_style(para, font_name, font_sizes["heading"],
                                        bold=True, align='left',
                                        indent=char_indent, line_spacing=line_spacing)
        elif ptype == 'heading3':
            font_name = fonts['heading3']
            self._apply_paragraph_style(para, font_name, font_sizes["heading"],
                                        bold=True, align='left',
                                        indent=char_indent, line_spacing=line_spacing)
        elif ptype == 'heading4':
            font_name = fonts['heading4']
            self._apply_paragraph_style(para, font_name, font_sizes["heading"],
                                        bold=False, align='left',
                                        indent=char_indent, line_spacing=line_spacing)
        else:
            font_name = fonts['body']
            self._apply_paragraph_style(para, font_name, font_sizes["body"],
                                        bold=False, align='justify',
                                        indent=char_indent, line_spacing=line_spacing)
        return para

    def _add_page_numbers(self, doc):
        section = doc.sections[0]
        try:
            doc.settings.odd_and_even_pages_header_footer = True
        except Exception:
            settings_el = doc.settings._element
            if settings_el.find(qn('w:evenAndOddHeaders')) is None:
                settings_el.append(OxmlElement('w:evenAndOddHeaders'))

        section.odd_and_even_pages_header_footer = True
        section.footer_distance = Cm(0.7)

        odd_footer = section.footer
        even_footer = section.even_page_footer
        odd_footer.is_linked_to_previous = False
        even_footer.is_linked_to_previous = False

        for para in odd_footer.paragraphs:
            para.clear()
        for para in even_footer.paragraphs:
            para.clear()

        def _build_footer_line(footer, align, pad_fullwidth):
            if footer.paragraphs:
                para = footer.paragraphs[0]
            else:
                para = footer.add_paragraph()

            para.alignment = align

            if pad_fullwidth:
                run0 = para.add_run("　")
                self._set_font(run0, 'Times New Roman', 14, bold=False)

            run1 = para.add_run("— ")
            self._set_font(run1, 'Times New Roman', 14, bold=False)

            run2 = para.add_run()
            self._set_font(run2, 'Times New Roman', 14, bold=False)
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run2._r.append(fldChar1)

            run3 = para.add_run()
            self._set_font(run3, 'Times New Roman', 14, bold=False)
            instrText = OxmlElement('w:instrText')
            instrText.text = 'PAGE'
            run3._r.append(instrText)

            run4 = para.add_run()
            self._set_font(run4, 'Times New Roman', 14, bold=False)
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            run4._r.append(fldChar2)

            run5 = para.add_run()
            self._set_font(run5, 'Times New Roman', 14, bold=False)
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')
            run5._r.append(fldChar3)

            run6 = para.add_run(" —")
            self._set_font(run6, 'Times New Roman', 14, bold=False)

            if not pad_fullwidth:
                run7 = para.add_run("　")
                self._set_font(run7, 'Times New Roman', 14, bold=False)

        _build_footer_line(odd_footer, WD_ALIGN_PARAGRAPH.RIGHT, pad_fullwidth=True)
        _build_footer_line(even_footer, WD_ALIGN_PARAGRAPH.LEFT, pad_fullwidth=False)

    def convert_file(self, input_path: str, output_path: str) -> bool:
        try:
            if self.progress_callback:
                self.progress_callback("正在读取文件...", 10)

            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()

            if self.progress_callback:
                self.progress_callback("正在创建文档...", 30)

            doc = Document()
            self._set_page_setup(doc)

            if self.progress_callback:
                self.progress_callback("正在解析Markdown...", 50)

            parsed = self._parse_markdown(content)

            if self.progress_callback:
                self.progress_callback("正在转换内容...", 70)

            prev_was_title = False

            for item in parsed:
                if len(item) == 4:
                    level, text, ptype, extra = item
                else:
                    level, text, ptype = item
                    extra = None

                if ptype == 'table' and extra:
                    self._create_table(doc, extra)
                    prev_was_title = False
                    continue

                if ptype == 'signature':
                    if extra == 'organization':
                        self._add_empty_paragraph(doc)
                        self._add_empty_paragraph(doc)
                    if extra == 'date':
                        self._add_empty_paragraph(doc)
                    self._create_signature_paragraph(doc, text, extra)
                    prev_was_title = False
                    continue

                if not text.strip():
                    continue

                if prev_was_title and ptype != 'title':
                    self._add_empty_paragraph(doc)

                self._create_styled_paragraph(doc, text, ptype)
                prev_was_title = (ptype == 'title')

            if self.progress_callback:
                self.progress_callback("正在添加页码...", 90)

            self._add_page_numbers(doc)

            if self.progress_callback:
                self.progress_callback("正在保存文件...", 95)

            doc.save(output_path)

            if self.progress_callback:
                self.progress_callback(f"转换完成: {Path(output_path).name}", 100)

            return True

        except Exception as e:
            if self.progress_callback:
                self.progress_callback(f"转换失败: {str(e)}", 0)
            return False

    def convert_text(self, markdown_text: str, output_path: str) -> bool:
        try:
            if self.progress_callback:
                self.progress_callback("正在创建文档...", 30)

            doc = Document()
            self._set_page_setup(doc)

            if self.progress_callback:
                self.progress_callback("正在解析Markdown...", 50)

            parsed = self._parse_markdown(markdown_text)

            if self.progress_callback:
                self.progress_callback("正在转换内容...", 70)

            prev_was_title = False

            for item in parsed:
                if len(item) == 4:
                    level, text, ptype, extra = item
                else:
                    level, text, ptype = item
                    extra = None

                if ptype == 'table' and extra:
                    self._create_table(doc, extra)
                    prev_was_title = False
                    continue

                if ptype == 'signature':
                    if extra == 'organization':
                        self._add_empty_paragraph(doc)
                        self._add_empty_paragraph(doc)
                    if extra == 'date':
                        self._add_empty_paragraph(doc)
                    self._create_signature_paragraph(doc, text, extra)
                    prev_was_title = False
                    continue

                if not text.strip():
                    continue

                if prev_was_title and ptype != 'title':
                    self._add_empty_paragraph(doc)

                self._create_styled_paragraph(doc, text, ptype)
                prev_was_title = (ptype == 'title')

            if self.progress_callback:
                self.progress_callback("正在添加页码...", 90)

            self._add_page_numbers(doc)

            if self.progress_callback:
                self.progress_callback("正在保存文件...", 95)

            doc.save(output_path)

            if self.progress_callback:
                self.progress_callback(f"转换完成: {Path(output_path).name}", 100)

            return True

        except Exception as e:
            if self.progress_callback:
                self.progress_callback(f"转换失败: {str(e)}", 0)
            return False
